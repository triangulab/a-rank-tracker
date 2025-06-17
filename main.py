import tkinter as tk
from tkinter import messagebox
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import RGBColor
import time
import undetected_chromedriver as uc

def launch_browser():
    options = uc.ChromeOptions()
    options.add_argument(r"--user-data-dir=C:/Users/flexa/AppData/Local/Google/Chrome/User Data")
    options.add_argument("--profile-directory=Profile 1")
    return uc.Chrome(options=options, headless=False)


def fetch_menu_items(driver, url):
    driver.get(url)
    time.sleep(5)

    print("[DEBUG] Starting smart scroll...")

    scroll_pause_time = 1.0
    scroll_attempts = 0
    max_scrolls = 50
    last_scroll_position = -1

    for _ in range(max_scrolls):
        driver.execute_script("window.scrollBy(0, 500);")
        time.sleep(scroll_pause_time)
        current_position = driver.execute_script("return window.pageYOffset;")
        if current_position == last_scroll_position:
            scroll_attempts += 1
            if scroll_attempts > 3:
                break
        else:
            scroll_attempts = 0
        last_scroll_position = current_position

    print("[DEBUG] Finished scrolling")

    # 🔻 NEW: Save what Selenium actually sees in the browser
    with open("page_dump.html", "w", encoding="utf-8") as f:
        f.write(driver.page_source)
    print("[DEBUG] Dumped live HTML to page_dump.html")

    # Try to find menu items
    containers = driver.find_elements(By.XPATH, "//div[@data-anchor-id='MenuItem']")
    print(f"[DEBUG] Found {len(containers)} menu items")

    items = []

    for c in containers:
        try:
            button = c.find_element(By.XPATH, ".//div[@role='button' and @aria-label]")
            aria = button.get_attribute("aria-label")
            print(f"[DEBUG] aria-label: {aria}")
        except:
            continue

        if "$" in aria:
            try:
                name_part, price_part = aria.rsplit("$", 1)
                name = name_part.strip()
                price = f"${price_part.strip()}"
            except:
                name = aria.strip()
                price = "Check Price"
        else:
            name = aria.strip()
            price = "Check Price"

        items.append({
            "category": "Menu",
            "name": name,
            "price": price,
            "description": ""
        })

    print(f"[DEBUG] Parsed {len(items)} menu items")
    return items


def create_word_doc(menu_items, filename="menu_output.docx"):
    doc = Document()
    current_category = None

    for item in menu_items:
        if item["category"] != current_category:
            current_category = item["category"]
            doc.add_heading(current_category, level=1)

        p = doc.add_paragraph()
        run = p.add_run(item["name"] + " ")
        run.bold = True

        run = p.add_run(item["price"])
        run.font.color.rgb = RGBColor(0, 102, 204)  # Blue

        if item["description"]:
            desc = doc.add_paragraph(item["description"])
            desc.runs[0].font.color.rgb = RGBColor(0, 153, 0)  # Green

    doc.save(filename)
    return filename

def download_menu():
    url = url_entry.get().strip()

    # Add https:// if missing
    if not url.startswith("http"):
        url = "https://" + url

    # Basic URL check
    if "doordash.com" not in url:
        messagebox.showerror("Error", "Please enter a valid DoorDash link.")
        return

    status_var.set("Loading browser...")
    root.update()

    try:
        driver = launch_browser()
        status_var.set("Scraping menu...")
        root.update()
        menu_items = fetch_menu_items(driver, url)
        status_var.set("Saving Word file...")
        root.update()
        doc_name = create_word_doc(menu_items)
        status_var.set("Done! Saved to " + doc_name)
        driver.quit()
    except Exception as e:
        status_var.set("Failed.")
        messagebox.showerror("Error", str(e))

# --- GUI ---
root = tk.Tk()
root.title("DoorDash Menu Extractor")

tk.Label(root, text="Enter DoorDash menu link:").pack(pady=(10, 0))

url_entry = tk.Entry(root, width=60)
url_entry.pack(pady=5)

tk.Button(root, text="Download Menu", command=download_menu).pack(pady=10)

status_var = tk.StringVar()
status_label = tk.Label(root, textvariable=status_var, fg="blue")
status_label.pack()

root.mainloop()
